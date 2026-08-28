import json
import sys
from pathlib import Path

from docx import Document


def text_of_paragraph(paragraph):
    return paragraph.text.strip()


def collect_tables(tables):
    result = []
    for table in tables:
        result.append(
            [[cell.text.strip().replace("\n", " | ") for cell in row.cells] for row in table.rows]
        )
    return result


def main(path_string):
    path = Path(path_string)
    document = Document(path)
    styles = []
    for style in document.styles:
        if style.type == 1:
            font = style.font
            styles.append({
                "name": style.name,
                "font": font.name,
                "size_pt": round(font.size.pt, 2) if font.size else None,
                "bold": font.bold,
                "italic": font.italic,
            })

    paragraphs = []
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = text_of_paragraph(paragraph)
        if text:
            paragraphs.append({
                "number": index,
                "style": paragraph.style.name,
                "alignment": str(paragraph.alignment),
                "text": text,
            })

    sections = []
    for section in document.sections:
        sections.append({
            "page_width_inches": round(section.page_width.inches, 2),
            "page_height_inches": round(section.page_height.inches, 2),
            "margins_inches": {
                "top": round(section.top_margin.inches, 2),
                "right": round(section.right_margin.inches, 2),
                "bottom": round(section.bottom_margin.inches, 2),
                "left": round(section.left_margin.inches, 2),
            },
            "header": [text_of_paragraph(p) for p in section.header.paragraphs if text_of_paragraph(p)],
            "footer": [text_of_paragraph(p) for p in section.footer.paragraphs if text_of_paragraph(p)],
        })

    output = {
        "file": str(path),
        "paragraphs": paragraphs,
        "tables": collect_tables(document.tables),
        "sections": sections,
        "inline_shapes": len(document.inline_shapes),
        "styles": styles,
    }
    mode = sys.argv[2] if len(sys.argv) > 2 else "full"
    if mode == "outline":
        outline = [
            entry for entry in paragraphs
            if "Heading" in entry["style"]
            or entry["style"] in {"Title", "Subtitle"}
            or entry["text"].upper().startswith(("CHAPTER", "TABLE OF CONTENTS", "ACKNOWLEDG", "ABSTRACT", "REFERENCES", "APPENDIX"))
        ]
        print(json.dumps({"file": str(path), "outline": outline, "table_count": len(output["tables"]), "sections": sections, "inline_shapes": len(document.inline_shapes)}, ensure_ascii=False, indent=2))
    elif mode == "text":
        for entry in paragraphs:
            print(f'{entry["number"]:03d} [{entry["style"]}] {entry["text"]}')
    else:
        print(json.dumps(output, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main(sys.argv[1])
