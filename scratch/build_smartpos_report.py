from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = r"C:\Users\Khun Thi Han\Documents\antigravity\kind-shannon\SMARTPOS_Internship_Project_Report.docx"


def shade(paragraph, fill="E7EEF8"):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def add_para(doc, text="", bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def figure_placeholder(doc, number, title, detail):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"Figure {number}. {title}")
    r.bold = True
    r.italic = True
    placeholder = doc.add_paragraph()
    placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    placeholder.paragraph_format.space_after = Pt(10)
    placeholder.paragraph_format.left_indent = Inches(0.35)
    placeholder.paragraph_format.right_indent = Inches(0.35)
    placeholder.add_run(f"[Insert {detail} here]").italic = True
    shade(placeholder, "F2F2F2")


def add_toc_line(doc, label, page):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.1), 2, 2)
    p.add_run(label)
    p.add_run("\t" + str(page))


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(0.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)

    for name, size, color in [("Heading 1", 16, "1F4E79"), ("Heading 2", 13, "1F4E79"), ("Heading 3", 11, "000000")]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if name == "Heading 1" else 8)
        style.paragraph_format.space_after = Pt(6)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("SMARTPOS Internship Project Report | Page ")
    page_number(footer)
    return doc


def main():
    doc = setup_document()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("SMARTPOS: MULTI-BRANCH POINT OF SALE\nAND INVENTORY MANAGEMENT SYSTEM")
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Times New Roman"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Internship Project Report").italic = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    p.add_run("Student Name: [To be completed]\nStudent ID / Class: [To be completed]\nSupervisor: [To be completed]\nSubmission Date: [To be completed]")
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("CONTENTS").bold = True
    p.runs[0].font.size = Pt(16)
    toc = [
        ("Acknowledgement", 1), ("Project Summary", 2), ("Motivation", 3),
        ("Objectives of the System", 3), ("Benefits of the System", 4),
        ("PROJECT DETAIL", 5), ("1. Developing Environment", 5),
        ("2. System Design", 8), ("3. Database Design", 12),
        ("4. System Functionalities", 16), ("5. User Interface Design", 22),
        ("6. Evaluation and Testing", 27), ("Knowledge and Technologies", 31),
        ("Limitations and Future Work", 33), ("Conclusion", 35), ("References", 36),
    ]
    for label, page in toc:
        add_toc_line(doc, label, page)
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("LIST OF FIGURES").bold = True
    p.runs[0].font.size = Pt(16)
    figures = [
        "Figure 2.1. SMARTPOS High-Level Architecture",
        "Figure 2.2. Role-Based Access and Branch Scope Flow",
        "Figure 2.3. POS Checkout and Inventory Audit Flow",
        "Figure 2.4. Sales Order Fulfillment and Delivery Flow",
        "Figure 3.1. SMARTPOS Entity-Relationship Diagram",
        "Figure 5.1. Sign-in Page", "Figure 5.2. Dashboard", "Figure 5.3. Sales Voucher / POS",
        "Figure 5.4. Payment Dialog", "Figure 5.5. Sales Orders", "Figure 5.6. Delivery Management",
        "Figure 5.7. Inventory", "Figure 5.8. Purchase Orders", "Figure 5.9. Outstanding Debts",
        "Figure 5.10. Staff Permission Management", "Figure 5.11. Reports", "Figure 5.12. Language Switcher",
    ]
    for index, label in enumerate(figures, start=1):
        add_toc_line(doc, label, index + 7)
    doc.add_page_break()

    add_heading(doc, "ACKNOWLEDGEMENT", 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "I would like to express my sincere gratitude to everyone who supported me throughout the completion of this internship project. Their guidance, encouragement, and practical advice made it possible to develop and document the SMARTPOS system.")
    add_para(doc, "I am especially grateful to my supervisor, [Supervisor Name], for providing direction, reviewing my work, and helping me strengthen both the technical implementation and the project report. I also thank the lecturers and staff of [University / Department] for the knowledge and support provided during my studies and internship period.")
    add_para(doc, "Finally, I would like to thank my family, friends, and classmates for their continual support and encouragement. Their confidence in me was an important source of motivation during this project.")

    add_heading(doc, "PROJECT SUMMARY", 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "SMARTPOS is a web-based Multi-Branch Point of Sale and Inventory Management System developed to help retail businesses manage daily sales, stock, staff access, purchasing, delivery operations, customer debts, and business reports from one integrated platform. The system is designed for businesses operating more than one branch and therefore places strong emphasis on branch-level data isolation and accurate inventory records.")
    add_para(doc, "The system provides three roles: Owner, Manager, and Cashier. Owners can access all branches and modules. Managers can manage permitted operations within their assigned branch. Cashiers are restricted to operational modules such as the Sales Voucher, delivery, and outstanding-debt collection. Each request is checked against role permissions and, for non-owners, the assigned branch boundary.")
    add_para(doc, "SMARTPOS supports direct POS checkout, sales orders, order fulfillment, delivery tracking, split payments, customer credit, purchase-order receiving, expense records, inventory adjustment, stock transfers, and management reports. Stock changes are paired with inventory-log records so that sales, purchases, adjustments, and transfers can be traced. The user interface supports English and Myanmar language modes to improve usability for local retail teams.")

    add_heading(doc, "MOTIVATION", 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Many small and medium retail businesses still depend on handwritten records, spreadsheet files, or disconnected applications. These methods can create inconsistent stock counts, delayed reporting, incorrect sales totals, difficulty following up customer debt, and weak control over which staff member can access sensitive information.")
    add_para(doc, "The main motivation for SMARTPOS is to provide a practical system that brings sales and stock data together while retaining appropriate control for a multi-branch business. The system should make it easy for staff to complete daily work, while giving managers and owners reliable information for decision-making.")

    add_heading(doc, "OBJECTIVES OF THE SYSTEM", 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    for item in [
        "To develop a centralized POS and inventory platform for multiple retail branches.",
        "To record sales, payments, stock movements, expenses, purchases, and outstanding debt accurately.",
        "To enforce role-based and branch-based access control for Owner, Manager, and Cashier users.",
        "To prevent invalid operations such as cross-branch changes, overpayment of debt, selling below cost, and stock deduction without an audit record.",
        "To provide bilingual English and Myanmar interfaces for everyday retail operations.",
        "To generate dashboards and reports that support timely business decisions.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "BENEFITS OF THE SYSTEM", 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    for item in [
        "Faster checkout through a focused Sales Voucher interface.",
        "Current branch-level stock visibility and low-stock awareness.",
        "Traceable inventory changes through inventory logs.",
        "Safer staff access using permissions and branch isolation.",
        "Better control of purchase receiving, moving-average cost, expenses, and customer debt.",
        "Accessible reporting for sales performance and operational monitoring.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "PROJECT DETAIL", 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_heading(doc, "1. Developing Environment", 1)
    add_para(doc, "SMARTPOS is a modern full-stack web application. It uses a component-based user interface, server-side API handlers, a relational database model, and automated test scripts. The following tools and technologies were used to develop the system.")
    tech = [
        ("Next.js 15", "Provides the App Router, page routing, server-side API route handlers, and production build process."),
        ("React 19", "Builds interactive dashboard, POS, form, dialog, and table components."),
        ("TypeScript", "Adds type checking for safer business logic, data models, and UI components."),
        ("Prisma ORM", "Connects application logic to the relational database and supports transactional operations."),
        ("MySQL", "Stores business data including staff, branches, products, stock levels, sales, purchases, and audit records."),
        ("Tailwind CSS and Radix UI", "Provide responsive styling and reusable accessible interface components."),
        ("Zustand", "Maintains client-side POS cart state."),
        ("VS Code, Git, npm, and tsx", "Support coding, version control, dependency management, and test execution."),
    ]
    for name, description in tech:
        add_para(doc, f"{name}: {description}", bold_prefix=f"{name}:")

    add_heading(doc, "2. System Design", 1)
    add_para(doc, "SMARTPOS uses a layered web architecture. Users interact with Next.js pages and React components. These pages request data from protected API routes. The API layer validates the staff session, checks permissions and branch scope, applies business rules, and uses Prisma to read or update the database. The database stores operational records and preserves inventory and transaction history.")
    figure_placeholder(doc, "2.1", "SMARTPOS High-Level Architecture", "architecture diagram showing User Interface → API Layer → Authentication/RBAC → Prisma → MySQL")
    add_heading(doc, "2.1 Role-Based Access and Branch Scope", 2)
    add_para(doc, "Every protected action obtains the current staff member from the session. The permission helper determines whether the requested module and action are allowed. Owner access is global. For Managers and Cashiers, any request that targets another branch is rejected. This rule prevents a staff member in one branch from reading or changing another branch's inventory, sales order, expense, or staff data.")
    figure_placeholder(doc, "2.2", "Role-Based Access and Branch Scope Flow", "flowchart for Owner, Manager, and Cashier permission decisions")
    add_heading(doc, "2.2 POS Checkout and Inventory Audit Flow", 2)
    add_para(doc, "At checkout, SMARTPOS validates the selected branch, items, quantities, payment method, discount amount, currency rule, available stock, and cost-price protection. A successful checkout creates a transaction and transaction items, deducts stock for each sold variant, and creates matching inventory-log entries in a database transaction. If any validation or stock update fails, the operation is rolled back rather than leaving partial data.")
    figure_placeholder(doc, "2.3", "POS Checkout and Inventory Audit Flow", "flowchart from cart validation through transaction, stock deduction, and inventory log")
    add_heading(doc, "2.3 Sales Order, Fulfillment, and Delivery Flow", 2)
    add_para(doc, "Sales orders allow the business to record requested products, deposits, payment status, customer details, and delivery information before completion. Fulfillment through the POS workflow checks remaining quantities and available stock, creates the fulfillment transaction, decreases stock, writes inventory logs, and updates fulfilled quantities. Delivery management then records delivery-service details, proof or receipt information where required, status, and any store-paid delivery fee expense.")
    figure_placeholder(doc, "2.4", "Sales Order Fulfillment and Delivery Flow", "flowchart from sales order to fulfillment, delivery, completion, and debt follow-up")

    add_heading(doc, "3. Database Design", 1)
    add_para(doc, "The SMARTPOS database is organized around branches, staff, products, stock, financial transactions, and order lifecycles. The design uses foreign-key relationships to ensure that operational data is connected to the branch, user, item, and document that created it. A compound uniqueness rule on branch and product variant prevents duplicate stock-level rows for the same item in the same branch.")
    figure_placeholder(doc, "3.1", "SMARTPOS Entity-Relationship Diagram", "ER diagram generated from the Prisma schema")
    models = [
        ("Branch and Staff", "Branch stores branch details. Staff belongs to a branch and contains role and permission data."),
        ("Category, Product, and ProductVariant", "Organize the product catalog, selling price, cost price, barcode, and variant-specific details."),
        ("StockLevel and InventoryLog", "Store current quantity by branch and variant, plus a durable history of every stock change and its reason."),
        ("Transaction and TransactionItem", "Store POS checkout headers, payment information, discounts, totals, and individual sold items."),
        ("SalesOrder, SalesOrderItem, and OrderPayment", "Support preorder quantities, deposit/payment status, fulfillment progress, delivery data, and debt collection."),
        ("Supplier, PurchaseOrder, and PurchaseItem", "Support suppliers, purchase-order lifecycle, received stock, amounts paid, and cost information."),
        ("Customer, Expense, ExchangeRate, and AuditLog", "Support customer identity and credit, operating expenses, rate history, and accountable user activity."),
    ]
    for name, description in models:
        add_para(doc, f"{name}: {description}", bold_prefix=f"{name}:")

    add_heading(doc, "4. System Functionalities", 1)
    add_heading(doc, "4.1 Owner Functions", 2)
    for item in ["Access all branches and all application modules.", "Manage branches, staff, permissions, products, categories, and system setup.", "Review sales, inventory, purchase, expense, delivery, and report data across the business."]:
        add_bullet(doc, item)
    add_heading(doc, "4.2 Manager Functions", 2)
    for item in ["Manage permitted branch operations such as inventory, sales orders, purchases, expenses, and reports.", "View and update Cashier permissions only within the assigned branch when authorized.", "Remain blocked from accessing or changing data in other branches."]:
        add_bullet(doc, item)
    add_heading(doc, "4.3 Cashier Functions", 2)
    for item in ["Use the Sales Voucher for direct retail checkout.", "Process assigned delivery tasks and record delivery status information.", "Review and collect permitted outstanding payments.", "Remain blocked from sensitive administration modules such as setup, reports, staff, purchases, and inventory management."]:
        add_bullet(doc, item)
    add_heading(doc, "4.4 Sales Voucher / POS", 2)
    add_para(doc, "The Sales Voucher screen displays products, variants, current branch stock, the active cart, customer details, discount controls, payment choices, and receipt information. It supports cash, card, QR, split, and debt-related workflows according to the selected transaction type. Validation ensures that quantities are positive, discounts are valid, stock is available, and the effective selling price does not violate configured cost-price rules.")
    add_heading(doc, "4.5 Inventory Management", 2)
    add_para(doc, "Inventory management displays product stock by branch and supports controlled stock adjustment and transfers. Each successful stock movement has a reason such as sale, adjustment, transfer in, transfer out, purchase received, or sales-order delivery/fulfillment, enabling later audit and reconciliation.")
    add_heading(doc, "4.6 Purchases and Moving-Average Cost", 2)
    add_para(doc, "Purchase orders record supplier, branch, items, cost, payment status, and receiving status. When goods are received, stock increases and item costs can be recalculated using moving-average cost logic. This helps the business maintain more realistic cost values as purchase prices change over time.")
    add_heading(doc, "4.7 Outstanding Debt and Customer Records", 2)
    add_para(doc, "Sales orders and wholesale/credit sales can leave a remaining balance. The Outstanding module calculates remaining debt from the order total, applicable delivery fee, and payments already made. A collection cannot exceed the remaining balance, helping prevent accidental overpayment and keeping customer payment records consistent.")
    add_heading(doc, "4.8 Reports, Dashboard, and Language Support", 2)
    add_para(doc, "The dashboard provides operational summaries and charts. Reports support review of sales and related records, while export functionality supports further analysis. The English/Myanmar language provider stores the selected language in local storage and changes supported interface labels without requiring a new login.")

    add_heading(doc, "5. User Interface Design", 1)
    add_para(doc, "The SMARTPOS user interface uses a dashboard layout with a role-aware sidebar, top header, language control, theme control, notifications, and user menu. Navigation entries are filtered according to the user's read permission, reducing clutter and avoiding exposure of unavailable modules.")
    for number, title, detail in [
        ("5.1", "Sign-in Page", "screenshot of the staff sign-in form"),
        ("5.2", "Dashboard", "screenshot of dashboard cards and charts"),
        ("5.3", "Sales Voucher / POS", "screenshot showing product grid, cart, and branch context"),
        ("5.4", "Payment Dialog", "screenshot showing payment validation and checkout options"),
        ("5.5", "Sales Orders", "screenshot of sales-order list and lifecycle controls"),
        ("5.6", "Delivery Management", "screenshot showing delivery status and fee details"),
        ("5.7", "Inventory", "screenshot showing stock levels and inventory logs"),
        ("5.8", "Purchase Orders", "screenshot showing purchase receiving workflow"),
        ("5.9", "Outstanding Debts", "screenshot showing debt balance and payment collection"),
        ("5.10", "Staff Permission Management", "screenshot showing module read/write controls"),
        ("5.11", "Reports", "screenshot showing reporting filters and results"),
        ("5.12", "Language Switcher", "English and Myanmar interface screenshots"),
    ]:
        figure_placeholder(doc, number, title, detail)

    add_heading(doc, "6. Evaluation and Testing", 1)
    add_para(doc, "The project includes automated unit and integration test scripts that exercise role boundaries, business lifecycles, inventory integrity, language switching, and stress scenarios. The following areas should be executed and documented again immediately before submission, because the final result must match the exact version of the source code submitted with this report.")
    checks = [
        ("Role-Based Access Control", "Verify that Owners have broad access, Managers are limited to their branches, and Cashiers cannot access restricted administration modules."),
        ("Branch Isolation", "Attempt reads and writes against another branch and confirm that unauthorized requests return a forbidden result."),
        ("POS Validation", "Check quantities, discounts, payment values, cost-price rules, MMK currency handling, and available-stock validation."),
        ("Inventory Integrity", "Verify that a stock decrease or increase is accompanied by the correct InventoryLog record and that the recorded quantity matches the expected movement."),
        ("Sales Orders and Fulfillment", "Check deposits, remaining quantity, fulfillment progress, transaction creation, and no duplicate stock deduction."),
        ("Delivery and Debt", "Check delivery completion requirements, delivery-fee accounting, remaining-debt calculation, and overpayment prevention."),
        ("Purchase Orders", "Check receiving flow, stock increase, and moving-average-cost behavior."),
        ("Language and Build", "Run language-switcher tests and the production build to ensure that the application compiles successfully."),
    ]
    for title, explanation in checks:
        add_para(doc, f"{title}: {explanation}", bold_prefix=f"{title}:")
    add_heading(doc, "Recommended Evidence for Submission", 2)
    for item in [
        "Attach terminal screenshots or a test-results appendix showing the final test commands and pass/fail results.",
        "State the date, commit/version, database environment, and any seed data used for final testing.",
        "Do not claim a test passed unless it was run against the submitted version of the project.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "KNOWLEDGE AND TECHNOLOGIES", 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "This project developed practical understanding of full-stack web application development. The work required designing a relational schema, building responsive interfaces, writing API routes, applying server-side validation, managing authentication and role permissions, handling transactions, and creating automated tests.")
    add_para(doc, "The project also demonstrated why inventory systems require careful treatment of concurrent operations and audit trails. A sales record alone is not sufficient: the associated stock movement, responsible staff member, branch, reason, and business document must remain traceable. This principle guided the use of StockLevel and InventoryLog records within transactional workflows.")

    add_heading(doc, "LIMITATIONS AND FUTURE WORK", 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    future = [
        ("Offline Capability", "Add offline-first POS operation and safe synchronization when internet connectivity returns."),
        ("Barcode Hardware", "Integrate dedicated barcode scanners, receipt printers, cash drawers, and label printers."),
        ("Accounting Integration", "Add general-ledger, tax, profit-and-loss, and bank-reconciliation features."),
        ("Notification Service", "Add email, SMS, Telegram, or in-app alerts for low stock, overdue debt, order changes, and failed delivery."),
        ("Advanced Reports", "Provide scheduled reports, richer filters, profit-margin dashboards, and export templates."),
        ("Mobile Experience", "Develop a mobile-optimized manager view or companion application."),
        ("Security Hardening", "Use secure password hashing, stronger session controls, rate limiting, multi-factor authentication, backups, monitoring, and formal security review before production deployment."),
    ]
    for title, description in future:
        add_para(doc, f"{title}: {description}", bold_prefix=f"{title}:")

    add_heading(doc, "CONCLUSION", 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "SMARTPOS was developed as a practical solution for a retail business that requires accurate sales, inventory, staff, and financial controls across multiple branches. It brings daily operational activities into one system while preserving branch boundaries and recording important stock movements.")
    add_para(doc, "The project demonstrates how a modern web application can support POS checkout, stock control, sales orders, deliveries, customer debt, purchases, expenses, and reporting. With final testing evidence, screenshots, and the student's cover details added, this report provides a complete academic record of the system and its implementation.")

    add_heading(doc, "REFERENCES", 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    refs = [
        "Next.js Documentation. https://nextjs.org/docs",
        "React Documentation. https://react.dev/",
        "Prisma Documentation. https://www.prisma.io/docs",
        "MySQL Documentation. https://dev.mysql.com/doc/",
        "TypeScript Documentation. https://www.typescriptlang.org/docs/",
        "Tailwind CSS Documentation. https://tailwindcss.com/docs",
        "Project source code and technical documentation: SMARTPOS repository, accessed for this report.",
    ]
    for ref in refs:
        add_bullet(doc, ref)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
